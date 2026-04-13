#!/usr/bin/env python3
"""Automated full-text acquisition for NMPA guidance documents.

CMDE/NMPA websites use Ruishu bot detection, so direct page scraping
is not feasible. This script uses Vertex AI Search (or Google CSE) to
discover docx/doc download URLs, then downloads, validates content, and
extracts text with embedded images.

Features:
    - .doc support via LibreOffice headless conversion
    - Image extraction: embedded images saved to docs/public/images/
    - Content validation: rejects catalog/list documents
    - Cross-matching: content not matching searched title is matched
      against other index entries or saved to _unmatched/

Phase 1 (--discover): Search for doc/docx URLs via Vertex AI Search
Phase 2 (--download): Download, convert .doc, validate, extract, cross-match
Phase 3 (--cleanup): Remove fulltext files with bad content
Phase 4 (--rematch): Re-match _unmatched/ files to index entries missing fulltext

Environment variables:
    VERTEX_AI_PROJECT_ID, VERTEX_AI_SEARCH_APP_ID, VERTEX_AI_SEARCH_API_KEY
    GOOGLE_SEARCH_API_KEY, GOOGLE_SEARCH_ENGINE_ID  (legacy fallback)

Usage:
    python scripts/fetch_nmpa_fulltext.py --stats
    python scripts/fetch_nmpa_fulltext.py --discover --category general
    python scripts/fetch_nmpa_fulltext.py --download --category general
    python scripts/fetch_nmpa_fulltext.py --cleanup --category general
    python scripts/fetch_nmpa_fulltext.py --rematch
    python scripts/fetch_nmpa_fulltext.py --rematch --dry-run
"""

import argparse
import hashlib
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from io import BytesIO
from pathlib import Path
from typing import Optional

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s: %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)

try:
    import requests
except ImportError:
    logger.error("requests not installed. pip install requests")
    sys.exit(1)

KNOWLEDGE_ROOT = Path(__file__).parent.parent
INDEX_PATH = KNOWLEDGE_ROOT / "nmpa" / "guidance" / "_index.json"
FULLTEXT_DIR = KNOWLEDGE_ROOT / "nmpa" / "guidance" / "fulltext"
UNMATCHED_DIR = KNOWLEDGE_ROOT / "nmpa" / "guidance" / "_unmatched"
IMAGES_DIR = KNOWLEDGE_ROOT / "docs" / "public" / "images" / "nmpa-guidance"
DISCOVERED_URLS_PATH = KNOWLEDGE_ROOT / "nmpa" / "guidance" / "_discovered_urls.json"

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/125.0.0.0 Safari/537.36"
)

CMDE_DOCX_PATTERN = re.compile(
    r'https?://(?:www\.)?cmde\.org\.cn/[^\s"\'<>]+\.docx?',
    re.IGNORECASE,
)
NMPA_DOCX_PATTERN = re.compile(
    r'https?://(?:www\.)?nmpa\.gov\.cn/[^\s"\'<>]+\.docx?',
    re.IGNORECASE,
)

_SPAM_URL_THRESHOLD = 20


def _build_spam_url_set() -> set[str]:
    """Identify URLs that appear 20+ times across discoveries -- generic/catalog files."""
    try:
        with open(DISCOVERED_URLS_PATH, "r", encoding="utf-8") as f:
            discovered = json.load(f)
    except FileNotFoundError:
        return set()

    from collections import Counter
    url_counter: Counter[str] = Counter()
    for disc in discovered.values():
        urls = disc.get("urls", [])
        if not urls:
            old = disc.get("url", "")
            if old:
                urls = [old]
        url_counter.update(urls)

    return {u for u, c in url_counter.items() if c >= _SPAM_URL_THRESHOLD}


def _normalize_doc_number(dn: str) -> str:
    """Normalize announcement number to group siblings from the same notice."""
    dn = re.sub(r'附件\d+', '', dn).strip()
    dn = re.sub(r'[，、,].*$', '', dn).strip()
    return dn


def _build_announcement_url_pool(
    entries: list[dict], discovered: dict, spam_urls: set[str]
) -> dict[str, set[str]]:
    """Pool URLs from all entries sharing the same announcement number.

    Returns a mapping of entry_id -> set of sibling URLs (excluding spam).
    """
    groups: dict[str, list[dict]] = {}
    for e in entries:
        dn = e.get("doc_number", "")
        if not dn:
            continue
        base = _normalize_doc_number(dn)
        if not base:
            continue
        groups.setdefault(base, []).append(e)

    pool: dict[str, set[str]] = {}
    for group_entries in groups.values():
        if len(group_entries) < 2:
            continue
        all_urls: set[str] = set()
        for e in group_entries:
            disc = discovered.get(e["id"], {})
            urls = disc.get("urls", [])
            if not urls:
                old = disc.get("url", "")
                if old:
                    urls = [old]
            all_urls.update(u for u in urls if u not in spam_urls)

        if all_urls:
            for e in group_entries:
                pool[e["id"]] = all_urls

    return pool


VERTEX_AI_PROJECT_ID = os.environ.get("VERTEX_AI_PROJECT_ID", "")
VERTEX_AI_SEARCH_APP_ID = os.environ.get("VERTEX_AI_SEARCH_APP_ID", "")
VERTEX_AI_SEARCH_API_KEY = os.environ.get("VERTEX_AI_SEARCH_API_KEY", "")

GOOGLE_SEARCH_API_KEY = os.environ.get("GOOGLE_SEARCH_API_KEY", "")
GOOGLE_SEARCH_ENGINE_ID = os.environ.get("GOOGLE_SEARCH_ENGINE_ID", "")

CN_SECTION_PATTERNS = [
    re.compile(r'^[一二三四五六七八九十]+[、.．]'),
    re.compile(r'^（[一二三四五六七八九十]+）'),
    re.compile(r'^第[一二三四五六七八九十百]+[章节条]'),
    re.compile(r'^\d+[、.．]\s'),
    re.compile(r'^\d+\.\d+[\s\S]'),
    re.compile(r'^（\d+）'),
    re.compile(r'^附[录件]\s*[一二三四五六七八九十A-Z\d]'),
]

_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

_LIBREOFFICE_BIN: Optional[str] = None


# ---------------------------------------------------------------------------
# LibreOffice .doc -> .docx conversion
# ---------------------------------------------------------------------------

def _detect_libreoffice() -> Optional[str]:
    """Find LibreOffice binary. Caches result."""
    global _LIBREOFFICE_BIN
    if _LIBREOFFICE_BIN is not None:
        return _LIBREOFFICE_BIN or None
    for name in ("libreoffice", "soffice"):
        path = shutil.which(name)
        if path:
            _LIBREOFFICE_BIN = path
            logger.info(f"LibreOffice found: {path}")
            return path
    if sys.platform == "darwin":
        mac_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        if os.path.isfile(mac_path):
            _LIBREOFFICE_BIN = mac_path
            logger.info(f"LibreOffice found: {mac_path}")
            return mac_path
    _LIBREOFFICE_BIN = ""
    logger.warning(
        "LibreOffice not found. .doc files cannot be processed. "
        "Install: brew install --cask libreoffice (macOS) or "
        "apt-get install libreoffice-writer-nogui (Linux)"
    )
    return None


def convert_doc_to_docx(doc_bytes: bytes) -> Optional[bytes]:
    """Convert .doc (OLE2) to .docx via LibreOffice headless."""
    lo_bin = _detect_libreoffice()
    if not lo_bin:
        return None

    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = Path(tmpdir) / "input.doc"
        doc_path.write_bytes(doc_bytes)

        try:
            result = subprocess.run(
                [lo_bin, "--headless", "--convert-to", "docx",
                 "--outdir", tmpdir, str(doc_path)],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode != 0:
                logger.warning(f"    LibreOffice conversion failed: {result.stderr[:200]}")
                return None
        except subprocess.TimeoutExpired:
            logger.warning("    LibreOffice conversion timed out")
            return None

        docx_path = Path(tmpdir) / "input.docx"
        if not docx_path.exists():
            for f in Path(tmpdir).glob("*.docx"):
                docx_path = f
                break
        if not docx_path.exists():
            logger.warning("    LibreOffice produced no .docx output")
            return None

        return docx_path.read_bytes()


def _is_doc_format(data: bytes) -> bool:
    """Check if bytes represent an OLE2 .doc file (not .docx)."""
    return data[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'


def _is_docx_format(data: bytes) -> bool:
    """Check if bytes represent a .docx (ZIP-based) file."""
    return data[:4] == b'PK\x03\x04'


# ---------------------------------------------------------------------------
# Image extraction from .docx
# ---------------------------------------------------------------------------

def extract_images_from_docx(docx_bytes: bytes, slug: str) -> dict[str, str]:
    """Extract embedded images from a docx and save to IMAGES_DIR/{slug}/.

    Returns a mapping from relationship ID (rId) to the markdown image path,
    e.g. {"rId5": "/images/nmpa-guidance/my-slug/image1.png"}.
    """
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except ImportError:
        return {}

    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception:
        return {}

    image_map: dict[str, str] = {}
    image_dir = IMAGES_DIR / slug
    img_counter = 0

    for rel_id, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue
        try:
            image_part = rel.target_part
            blob = image_part.blob
            if not blob or len(blob) < 100:
                continue

            content_type = image_part.content_type or ""
            if "png" in content_type:
                ext = "png"
            elif "jpeg" in content_type or "jpg" in content_type:
                ext = "jpg"
            elif "gif" in content_type:
                ext = "gif"
            elif "svg" in content_type:
                ext = "svg"
            elif "emf" in content_type or "wmf" in content_type:
                continue
            else:
                ext = "png"

            img_counter += 1
            fname = f"image{img_counter}.{ext}"
            rel_path = f"/images/nmpa-guidance/{slug}/{fname}"

            image_dir.mkdir(parents=True, exist_ok=True)
            (image_dir / fname).write_bytes(blob)

            image_map[rel_id] = rel_path
        except Exception:
            continue

    if img_counter > 0:
        logger.info(f"    Extracted {img_counter} images to {image_dir.relative_to(KNOWLEDGE_ROOT)}")

    return image_map


# ---------------------------------------------------------------------------
# Web Search (Vertex AI Search / Google CSE)
# ---------------------------------------------------------------------------

class WebSearcher:
    """Unified search via Vertex AI Search (preferred) or Google CSE (legacy)."""

    VERTEX_URL = (
        "https://discoveryengine.googleapis.com/v1/projects/{project_id}"
        "/locations/global/collections/default_collection"
        "/engines/{app_id}/servingConfigs/default_search:searchLite"
    )
    GOOGLE_CSE_URL = "https://www.googleapis.com/customsearch/v1"

    def __init__(self):
        self.use_vertex = bool(
            VERTEX_AI_PROJECT_ID and VERTEX_AI_SEARCH_APP_ID
            and VERTEX_AI_SEARCH_API_KEY
        )
        self.use_google = bool(GOOGLE_SEARCH_API_KEY and GOOGLE_SEARCH_ENGINE_ID)
        if self.use_vertex:
            logger.info("Using Vertex AI Search (searchLite)")
        elif self.use_google:
            logger.info("Using Google CSE (legacy)")
        else:
            logger.warning(
                "No search API configured. Set VERTEX_AI_* or GOOGLE_SEARCH_* env vars."
            )

    @property
    def available(self) -> bool:
        return self.use_vertex or self.use_google

    def search(self, query: str, num: int = 10) -> list[dict]:
        if self.use_vertex:
            return self._vertex_search(query, num)
        if self.use_google:
            return self._google_search(query, num)
        return []

    def _vertex_search(self, query: str, num: int) -> list[dict]:
        query = re.sub(r"site:\S+\s*", "", query).strip()
        query = re.sub(r"filetype:\S+\s*", "", query).strip()
        url = self.VERTEX_URL.format(
            project_id=VERTEX_AI_PROJECT_ID,
            app_id=VERTEX_AI_SEARCH_APP_ID,
        )
        try:
            resp = requests.post(
                url,
                params={"key": VERTEX_AI_SEARCH_API_KEY},
                json={"query": query, "pageSize": min(num, 25)},
                headers={"Content-Type": "application/json"},
                timeout=20,
            )
            if resp.status_code != 200:
                logger.warning(f"Vertex AI Search HTTP {resp.status_code}: "
                               f"{resp.text[:200]}")
                return []
            results = []
            for r in resp.json().get("results", []):
                doc = r.get("document", {})
                derived = doc.get("derivedStructData", {})
                link = derived.get("link", "")
                title = derived.get("title", "")
                snippet = ""
                for s in derived.get("snippets", []):
                    if s.get("snippet"):
                        snippet = re.sub(r"<[^>]+>", "", s["snippet"])
                        break
                if title or link:
                    results.append({"title": title, "link": link,
                                    "snippet": snippet[:500]})
            return results
        except Exception as e:
            logger.warning(f"Vertex AI Search error: {e}")
            return []

    def _google_search(self, query: str, num: int) -> list[dict]:
        try:
            params = {
                "key": GOOGLE_SEARCH_API_KEY,
                "cx": GOOGLE_SEARCH_ENGINE_ID,
                "q": query,
                "num": min(num, 10),
            }
            resp = requests.get(self.GOOGLE_CSE_URL, params=params, timeout=15)
            resp.raise_for_status()
            results = []
            for item in resp.json().get("items", []):
                results.append({
                    "title": item.get("title", ""),
                    "link": item.get("link", ""),
                    "snippet": item.get("snippet", "")[:300],
                })
            return results
        except Exception as e:
            logger.warning(f"Google CSE error: {e}")
            return []


# ---------------------------------------------------------------------------
# Index & URL management
# ---------------------------------------------------------------------------

def load_index() -> dict:
    with open(INDEX_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def save_index(data: dict):
    with open(INDEX_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    logger.info(f"Saved index: {INDEX_PATH}")


def load_discovered_urls() -> dict:
    if DISCOVERED_URLS_PATH.exists():
        with open(DISCOVERED_URLS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_discovered_urls(data: dict):
    with open(DISCOVERED_URLS_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def generate_slug(title_zh: str, doc_number: str) -> str:
    dn_match = re.search(r'(\d{4})\D*?(\d+)', doc_number)
    if dn_match:
        return f"cmde-{dn_match.group(1)}-{dn_match.group(2)}"
    short_hash = hashlib.md5(title_zh.encode()).hexdigest()[:8]
    return f"cmde-{short_hash}"


def fulltext_exists(entry: dict) -> bool:
    slug = entry.get("slug", "")
    eid = entry.get("id", "")
    for name in [slug, re.sub(r'[^\w\-.]', '_', eid)]:
        if not name:
            continue
        if (FULLTEXT_DIR / f"{name}.md").exists():
            return True
        if (FULLTEXT_DIR / f"{name}.zh.md").exists():
            return True
    return False


# ---------------------------------------------------------------------------
# Content Validation
# ---------------------------------------------------------------------------

def is_catalog_content(text: str) -> bool:
    """Detect if extracted text is a catalog/list of guidance titles."""
    if not text:
        return True

    text_kb = max(len(text) / 1024, 0.1)
    guideline_count = text.count("指导原则")
    density = guideline_count / text_kb

    if density > 8 and guideline_count > 20:
        return True

    attachment_list = re.findall(r'附件\s*\d+[：:.]\s*.*指导原则', text)
    if len(attachment_list) > 5:
        return True

    lines = text.split('\n')
    non_empty_lines = [l for l in lines if l.strip()]
    if non_empty_lines and len(non_empty_lines) < 200:
        title_like_lines = sum(
            1 for l in non_empty_lines
            if '指导原则' in l and len(l.strip()) < 80
        )
        if title_like_lines > 15 and title_like_lines / len(non_empty_lines) > 0.4:
            return True

    return False


def _title_keywords(title: str) -> set[str]:
    """Extract meaningful keywords from a Chinese title."""
    title_core = re.sub(r'[（()）\[\]【】（）/／]', '', title)
    title_core = re.sub(r'\d{4}年.*?修订版', '', title_core)
    title_core = re.sub(r'第\d+号', '', title_core)
    generic_terms = {
        '指导原则', '注册', '审查', '技术', '医疗器械', '体外诊断',
        '试剂', '产品', '临床', '研究', '评价', '评估', '注册审查',
        '技术审查', '注册技术审查',
    }
    keywords = set()
    for seg in re.split(r'[、，,\s]+', title_core):
        seg = seg.strip()
        if not seg:
            continue
        if seg in generic_terms:
            continue
        if len(seg) >= 2:
            keywords.add(seg)
        if len(seg) > 6:
            for sub_term in generic_terms:
                if sub_term in seg:
                    parts = seg.split(sub_term)
                    for p in parts:
                        p = p.strip()
                        if len(p) >= 2 and p not in generic_terms:
                            keywords.add(p)
    return keywords


def _normalize_title(title: str) -> str:
    """Normalize a Chinese title for comparison (remove spaces, punct variants)."""
    t = re.sub(r'\s+', '', title)
    t = t.replace('/', '').replace('（', '(').replace('）', ')')
    t = re.sub(r'\(\d{4}年.*?\)', '', t)
    return t


def _extract_core_subject(title: str) -> str:
    """Extract the core product/test name by stripping generic suffixes."""
    s = title.strip()
    s = s.replace('\uff08', '(').replace('\uff09', ')')
    s = re.sub(r'\(\d{4}\u5e74.*?\)', '', s)
    s = re.sub(
        r'(\u6ce8\u518c|\u6280\u672f|\u5ba1\u67e5|\u5ba1\u8bc4|\u6307\u5bfc'
        r'|\u539f\u5219|\u8981\u70b9|\u6307\u5357|\u8bd5\u884c'
        r'|\u4fee\u8ba2\u7248?|\u4ea7\u54c1|\u901a\u7528|\u89c4\u8303)+', '', s)
    s = re.sub(r'\u7b2c\d+\u53f7.*$', '', s)
    s = re.sub(r'\u7b2c\d+\u90e8\u5206.*$', '', s)
    s = re.sub(r'[（()）\[\]【】/／：:\s]', '', s)
    return s.strip()


def _check_attachment_body_mismatch(text: str, expected_title: str) -> bool:
    """Detect if text has the expected title in header but body content
    is actually about a different guidance document (NMPA attachment pattern).

    Returns True if the body is about a DIFFERENT topic than expected_title.
    """
    lines = text.split('\n')

    attachment_markers = (
        '## \u9644\u4ef6', '\u9644\u4ef6', '# \u9644\u4ef6',
        '**\u9644\u4ef6**', '## \u9644\u4ef6\uff1a', '\u9644\u4ef6\uff1a', '\u9644\u4ef6:',
    )

    attachment_idx = -1
    for i, line in enumerate(lines[:30]):
        stripped = line.strip()
        if stripped in attachment_markers or stripped.startswith('## \u9644\u4ef6'):
            attachment_idx = i
            break

    if attachment_idx < 0:
        return False

    for j in range(attachment_idx + 1, min(attachment_idx + 8, len(lines))):
        line = lines[j].strip()
        line = re.sub(r'^#+\s*', '', line).strip()
        line = re.sub(r'^\*+\s*', '', line).rstrip('*').strip()
        if not line:
            continue
        if re.search(r'[\u4e00-\u9fff]', line) and len(line) >= 4:
            body_subject = line[:100]
            expected_core = _extract_core_subject(expected_title)
            body_core = _extract_core_subject(body_subject)

            if not expected_core or not body_core or len(expected_core) < 2:
                return False

            if expected_core in body_core or body_core in expected_core:
                return False

            title_sim = _title_similarity(body_subject, expected_title)
            if title_sim >= 0.5:
                return False

            if len(expected_core) >= 4 and len(body_core) >= 4:
                exp_bg = set(expected_core[i:i+2]
                             for i in range(len(expected_core)-1))
                body_bg = set(body_core[i:i+2]
                              for i in range(len(body_core)-1))
                if exp_bg and body_bg:
                    overlap = exp_bg & body_bg
                    score = len(overlap) / max(len(exp_bg), len(body_bg))
                    if score < 0.3:
                        return True

            return True

    return False


def content_matches_title(text: str, expected_title: str) -> bool:
    """Check if the document content plausibly matches the expected title.

    Also detects the NMPA "attachment mismatch" pattern where the title
    appears in the header but the body is a different guidance document.
    """
    if not text or not expected_title:
        return False

    first_section = text[:5000]
    first_flat = re.sub(r'\s*\n#+\s*', '', first_section[:3000])

    title_found = False
    if expected_title in first_section or expected_title in first_flat:
        title_found = True
    else:
        norm_expected = _normalize_title(expected_title)
        norm_first = _normalize_title(first_flat)
        if norm_expected in norm_first:
            title_found = True

    if title_found:
        if _check_attachment_body_mismatch(text, expected_title):
            return False
        return True

    title_keywords = _title_keywords(expected_title)
    if not title_keywords:
        return True

    keyword_hits = sum(1 for kw in title_keywords if kw in first_section)
    return keyword_hits / len(title_keywords) >= 0.3


def is_draft_content(text: str) -> bool:
    """Detect if extracted text is a draft for comments (征求意见稿).

    Only official/final versions should be included in the knowledge base.
    Draft documents are identified by explicit markers in the first ~3000 chars.
    """
    if not text:
        return False

    head = text[:3000]

    draft_patterns = [
        r'征求意见稿',
        r'征求\s*意见\s*稿',
    ]

    for pat in draft_patterns:
        if re.search(pat, head):
            return True

    return False


def validate_extracted_content(
    markdown: str, expected_title: str
) -> tuple[bool, str]:
    """Validate that extracted content is the actual guidance document.

    Returns (is_valid, reason).
    """
    if not markdown or len(markdown) < 200:
        return False, "too_short"

    if is_catalog_content(markdown):
        return False, "catalog"

    if is_draft_content(markdown):
        return False, "draft"

    if not content_matches_title(markdown, expected_title):
        return False, "title_mismatch"

    return True, "ok"


# ---------------------------------------------------------------------------
# Cross-matching: match mismatched content to other index entries
# ---------------------------------------------------------------------------

_BODY_START_RE = re.compile(
    r'^(本指导原则|本文|为进一步|为规范|为统一|为指导|根据|按照'
    r'|分析性能评估资料|本文涉及|本技术|本标准|本导则'
    r'|本指南)',
)

_ATTACHMENT_LINE_RE = re.compile(
    r'^(附件\s*[:：]?\s*\d*\s*$|## 附件|# 附件)',
    re.IGNORECASE,
)

_REVISION_SUFFIX_RE = re.compile(
    r'^[（(]\s*\d{4}\s*年?\s*(修订版|第.+版)\s*[）)]\s*$',
)


def extract_actual_title(markdown: str) -> str:
    """Extract the guidance document title from markdown content.

    Handles NMPA guidance document patterns:
    - "附件N" prefix lines (skipped)
    - Title split across multiple lines by hard returns
    - "(2024年修订版)" suffixes appended to title
    - Markdown heading markers (# / ##)
    - Bold markers (**title**)
    """
    lines = []
    for raw in markdown.split('\n'):
        stripped = raw.strip()
        if not stripped:
            continue
        if stripped.startswith('<!--'):
            continue
        if stripped.startswith('!['):
            continue
        clean = re.sub(r'^#+\s*', '', stripped).strip()
        clean = re.sub(r'^\*+\s*', '', clean).rstrip('*').strip()
        if not clean or clean == '---':
            continue
        lines.append(clean)

    if not lines:
        return ""

    title_parts: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]
        if _ATTACHMENT_LINE_RE.match(line) or _ATTACHMENT_LINE_RE.match(lines[i] if not line.startswith('#') else f'# {line}'):
            i += 1
            continue
        if re.match(r'^附件\s*[:：]?\s*$', line) or re.match(r'^附件\s*\d+\s*$', line):
            i += 1
            continue
        break

    while i < len(lines):
        line = lines[i]

        if _BODY_START_RE.match(line):
            break

        if re.match(r'^(Source:|Published:|---)', line):
            break
        if any(line.startswith(p) for p in _SKIP_LINE_PREFIXES):
            break

        if re.match(r'^指导原则编号', line):
            i += 1
            continue
        if re.match(r'^(二O|20\d{2}年\d+月)', line):
            i += 1
            continue

        if _REVISION_SUFFIX_RE.match(line):
            title_parts.append(line)
            i += 1
            break

        if len(line) > 60 and title_parts:
            break

        title_parts.append(line)
        i += 1

        cur_title = ''.join(title_parts)
        if re.search(r'(指导原则|指导文件|审查要求|技术指导|基本要求|技术要求)', cur_title):
            if i < len(lines) and not _REVISION_SUFFIX_RE.match(lines[i]):
                break

        if len(title_parts) >= 3:
            break

    if not title_parts:
        for line in lines:
            if re.search(r'[\u4e00-\u9fff]', line) and len(line) >= 4:
                return line[:100]
        return ""

    title = ''.join(title_parts)
    title = re.sub(r'\s+', '', title)
    return title[:150]


def _title_similarity(title_a: str, title_b: str) -> float:
    """Compute similarity between two Chinese titles (0.0 - 1.0)."""
    if not title_a or not title_b:
        return 0.0

    if title_a in title_b or title_b in title_a:
        return 0.9

    kw_a = _title_keywords(title_a)
    kw_b = _title_keywords(title_b)
    if not kw_a or not kw_b:
        return 0.0

    overlap = kw_a & kw_b
    union = kw_a | kw_b
    return len(overlap) / len(union) if union else 0.0


def cross_match_content(
    markdown: str, entries: list[dict], source_url: str = ""
) -> Optional[dict]:
    """Try to match extracted content against any index entry without fulltext.

    Returns the best matching entry, or None.
    """
    actual_title = extract_actual_title(markdown)
    if not actual_title:
        return None

    if is_catalog_content(markdown):
        return None

    best_entry = None
    best_score = 0.0

    for entry in entries:
        if fulltext_exists(entry):
            continue
        entry_title = entry.get("title", {}).get("zh", "")
        if not entry_title:
            continue

        score = _title_similarity(actual_title, entry_title)
        if score > best_score:
            best_score = score
            best_entry = entry

    if best_score >= 0.5 and best_entry:
        if content_matches_title(markdown, best_entry["title"]["zh"]):
            return best_entry

    return None


def save_unmatched(markdown: str, source_url: str, actual_title: str):
    """Save content that couldn't be matched to any index entry."""
    UNMATCHED_DIR.mkdir(parents=True, exist_ok=True)
    slug = hashlib.md5(actual_title.encode()).hexdigest()[:12]
    fp = UNMATCHED_DIR / f"{slug}.md"

    header = f"<!-- unmatched -->\n"
    header += f"<!-- actual_title: {actual_title} -->\n"
    header += f"<!-- source_url: {source_url} -->\n\n"

    fp.write_text(header + markdown, encoding="utf-8")
    logger.info(f"    Saved to _unmatched/{slug}.md ({actual_title[:40]})")


_SKIP_LINE_PREFIXES = (
    'Source:', 'Published:', '**Source',
    '\u98df\u836f\u76d1', '\u56fd\u5bb6\u98df\u54c1',
    '\u56fd\u5bb6\u836f\u54c1\u76d1\u7763\u7ba1\u7406\u5c40',
    '\u5173\u4e8e\u53d1\u5e03', '\u5173\u4e8e\u5370\u53d1',
)

_SKIP_LINE_EXACT = {'\u9644\u4ef6', '\u9644\u5f55'}


def _is_boilerplate_line(clean: str) -> bool:
    """Return True if line is boilerplate that should not be used as a title."""
    if not clean or clean == '---':
        return True
    if any(clean.startswith(p) for p in _SKIP_LINE_PREFIXES):
        return True
    if clean in _SKIP_LINE_EXACT:
        return True
    if re.match(r'^\u9644\u4ef6\s*\d*$', clean):
        return True
    return False


def _extract_unmatched_title(text: str) -> str:
    """Extract the actual guidance title from an _unmatched/ file,
    skipping boilerplate headers and metadata comments."""
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('<!--'):
            continue
        clean = re.sub(r'^#+\s*', '', line).strip()
        clean = re.sub(r'^\*+\s*', '', clean).rstrip('*').strip()
        if _is_boilerplate_line(clean):
            continue
        if re.search(r'[\u4e00-\u9fff]', clean) and len(clean) >= 4:
            return clean[:100]
    return ""


def _extract_candidate_titles(text: str, max_titles: int = 5) -> list[str]:
    """Extract multiple candidate title strings from content,
    skipping boilerplate. Used for deeper rematch attempts."""
    candidates = []
    seen: set[str] = set()
    for line in text.split('\n')[:50]:
        line = line.strip()
        if not line or line.startswith('<!--'):
            continue
        clean = re.sub(r'^#+\s*', '', line).strip()
        clean = re.sub(r'^\*+\s*', '', clean).rstrip('*').strip()
        if _is_boilerplate_line(clean):
            continue
        if re.search(r'[\u4e00-\u9fff]', clean) and len(clean) >= 4:
            title = clean[:100]
            if title not in seen:
                seen.add(title)
                candidates.append(title)
            if len(candidates) >= max_titles:
                break
    return candidates


def _extract_unmatched_body(text: str) -> str:
    """Extract body content from an _unmatched/ file (skip comment headers)."""
    lines = text.split('\n')
    body_lines = []
    past_comments = False
    for line in lines:
        if not past_comments:
            stripped = line.strip()
            if stripped.startswith('<!--') and stripped.endswith('-->'):
                continue
            if not stripped:
                continue
            past_comments = True
        body_lines.append(line)
    return '\n'.join(body_lines)


def rematch_unmatched(dry_run: bool = False):
    """Scan _unmatched/ files, try to match them to index entries missing fulltext.

    Moves matched files from _unmatched/ to fulltext/ with proper slug.
    """
    if not UNMATCHED_DIR.exists():
        logger.info("No _unmatched/ directory found")
        return

    index_data = load_index()
    entries = index_data.get("entries", [])

    unmatched_files = sorted(UNMATCHED_DIR.glob("*.md"))
    logger.info(f"Scanning {len(unmatched_files)} unmatched files...")

    matched = 0
    skipped_catalog = 0
    skipped_draft = 0
    no_match = 0
    index_modified = False

    for fp in unmatched_files:
        text = fp.read_text(encoding='utf-8', errors='replace')
        body = _extract_unmatched_body(text)

        if not body or len(body) < 200:
            continue

        if is_catalog_content(body):
            skipped_catalog += 1
            continue

        if is_draft_content(body):
            skipped_draft += 1
            continue

        actual_title = extract_actual_title(body)
        if not actual_title or len(actual_title) < 4:
            no_match += 1
            continue

        source_url = ""
        m = re.search(r'<!--\s*source_url:\s*(.*?)\s*-->', text)
        if m:
            source_url = m.group(1).strip()

        found_entry = None
        found_score = 0.0

        best_entry = None
        best_score = 0.0

        for entry in entries:
            if fulltext_exists(entry):
                continue
            entry_title = entry.get("title", {}).get("zh", "")
            if not entry_title:
                continue

            score = _title_similarity(actual_title, entry_title)
            if score > best_score:
                best_score = score
                best_entry = entry

        if best_score >= 0.5 and best_entry:
            entry_title = best_entry["title"]["zh"]
            if content_matches_title(body, entry_title):
                found_entry = best_entry
                found_score = best_score

        if not found_entry:
            no_match += 1
            continue

        entry_title = found_entry["title"]["zh"]
        logger.info(
            f"  REMATCH [{found_score:.2f}]: {fp.name}"
            f" -> {entry_title[:50]}"
        )

        if dry_run:
            matched += 1
            continue

        clean_body = re.sub(r'<!--\s*unmatched\s*-->\s*\n?', '', text)
        clean_body = re.sub(r'<!--\s*actual_title:.*?-->\s*\n?', '', clean_body)
        clean_body = re.sub(r'<!--\s*source_url:.*?-->\s*\n?', '', clean_body)
        clean_body = clean_body.lstrip('\n')

        if _save_fulltext(found_entry, clean_body, source_url, entries):
            index_modified = True

        fp.unlink()
        matched += 1
        logger.info(f"    Moved to fulltext, removed from _unmatched/")

    if index_modified and not dry_run:
        save_index(index_data)

    logger.info(
        f"\nRematch: {matched} matched"
        f", {skipped_catalog} catalogs skipped"
        f", {skipped_draft} drafts skipped"
        f", {no_match} no match"
        f"{' (dry-run)' if dry_run else ''}"
    )


# ---------------------------------------------------------------------------
# Phase 1: Discover docx URLs via web search
# ---------------------------------------------------------------------------

def extract_all_docx_urls(results: list[dict]) -> list[str]:
    """Extract ALL doc/docx download URLs from search results (deduplicated)."""
    urls = []
    seen = set()
    for result in results:
        link = result.get("link", "")
        snippet = result.get("snippet", "")
        for source in [link, snippet]:
            for pattern in [CMDE_DOCX_PATTERN, NMPA_DOCX_PATTERN]:
                for m in pattern.finditer(source):
                    url = m.group(0)
                    if url not in seen:
                        seen.add(url)
                        urls.append(url)
        if re.search(r'\.docx?$', link, re.IGNORECASE) and link not in seen:
            seen.add(link)
            urls.append(link)
    return urls


def discover_urls(searcher: WebSearcher, limit: int = 0, dry_run: bool = False,
                  category: str = ""):
    """Phase 1: Search for doc/docx URLs for entries without fulltext."""
    if not searcher.available:
        logger.error("No search API available. Cannot discover URLs.")
        return

    index_data = load_index()
    entries = index_data.get("entries", [])
    discovered = load_discovered_urls()

    candidates = [
        e for e in entries
        if not fulltext_exists(e) and e["id"] not in discovered
    ]
    if category:
        candidates = [e for e in candidates if e.get("category") == category]
    if limit > 0:
        candidates = candidates[:limit]

    logger.info(
        f"Discovering URLs for {len(candidates)} entries "
        f"({len(entries)} total, {len(discovered)} already searched)"
    )

    found = 0
    for i, entry in enumerate(candidates):
        title_zh = entry.get("title", {}).get("zh", "")
        eid = entry["id"]

        logger.info(f"[{i+1}/{len(candidates)}] {title_zh[:50]}...")

        if dry_run:
            continue

        all_urls: list[str] = []
        seen_urls: set[str] = set()

        short_title = re.sub(r'[（(].*?[）)]', '', title_zh).strip()

        queries = [
            f'{title_zh} 全文 docx',
            f'{short_title} 注册审查 指导原则',
            f'{title_zh}',
        ]

        page_urls: list[str] = []

        for query in queries:
            results = searcher.search(query, num=10)
            for url in extract_all_docx_urls(results):
                if url not in seen_urls:
                    seen_urls.add(url)
                    all_urls.append(url)
            for r in results:
                link = r.get("link", "")
                if link and not re.search(r'\.docx?$', link, re.IGNORECASE):
                    page_urls.append(link)
            time.sleep(0.3)

        if all_urls:
            discovered[eid] = {
                "urls": all_urls,
                "format": "docx",
            }
            logger.info(f"  FOUND {len(all_urls)} candidate URLs")
            found += 1
        else:
            discovered[eid] = {
                "urls": [],
                "page_url": page_urls[0] if page_urls else "",
                "format": "",
            }
            logger.info(f"  NOT FOUND")

        if (i + 1) % 10 == 0:
            save_discovered_urls(discovered)

        time.sleep(1)

    if not dry_run:
        save_discovered_urls(discovered)
    logger.info(f"\nDiscovery: {found} entries with URLs out of {len(candidates)}")


# ---------------------------------------------------------------------------
# Phase 2: Download, Convert, Extract, Validate, Cross-match
# ---------------------------------------------------------------------------

def download_file(url: str) -> Optional[bytes]:
    try:
        resp = requests.get(
            url,
            headers={"User-Agent": USER_AGENT},
            timeout=60,
            allow_redirects=True,
        )
        resp.raise_for_status()
        ct = resp.headers.get("content-type", "").lower()
        if len(resp.content) < 500 and "html" in ct:
            logger.warning(f"    Got HTML instead of document (blocked?)")
            return None
        return resp.content
    except Exception as e:
        logger.warning(f"    Download failed: {e}")
        return None


def file_to_docx_bytes(data: bytes, url: str) -> Optional[bytes]:
    """Ensure we have .docx bytes. Converts .doc via LibreOffice if needed."""
    if _is_docx_format(data):
        return data

    if _is_doc_format(data):
        logger.info(f"    .doc format detected, converting via LibreOffice...")
        converted = convert_doc_to_docx(data)
        if converted:
            logger.info(f"    Converted to .docx ({len(converted)/1024:.1f} KB)")
        return converted

    if url.lower().endswith('.doc'):
        logger.info(f"    Trying LibreOffice conversion for ambiguous format...")
        return convert_doc_to_docx(data)

    logger.info(f"    Skip: unrecognized file format")
    return None


def extract_docx_to_markdown(docx_bytes: bytes, image_map: dict[str, str] | None = None) -> str:
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. pip install python-docx")
        return ""

    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception as e:
        logger.error(f"    Failed to open docx: {e}")
        return ""

    lines: list[str] = []
    table_set: set[int] = set()

    body_el = doc.element.body
    for child in body_el:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue

            text = para.text.strip()

            has_images = False
            if image_map:
                for run in para.runs:
                    drawing_els = run._element.findall(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
                    ) or run._element.findall(
                        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
                    )
                    if drawing_els:
                        has_images = True

            if not text and not has_images:
                lines.append("")
                continue

            if text:
                style_name = para.style.name if para.style else ""
                md_line = _para_to_markdown(text, style_name, para)
                if md_line.startswith("#"):
                    lines.append("")
                    lines.append(md_line)
                    lines.append("")
                else:
                    if lines and lines[-1] and not lines[-1].startswith("#"):
                        lines.append("")
                    lines.append(md_line)

            if has_images and image_map:
                for rel_id, img_path in image_map.items():
                    blip_els = para._element.findall(
                        f'.//{{{_nsmap("a")}}}blip[@{{{_nsmap("r")}}}embed="{rel_id}"]'
                    )
                    if blip_els:
                        lines.append(f"\n![](/images/nmpa-guidance/{img_path.split('/images/nmpa-guidance/')[-1]})\n")

        elif tag == "tbl":
            tbl_id = id(child)
            if tbl_id in table_set:
                continue
            table_set.add(tbl_id)
            for tbl in doc.tables:
                if tbl._element is child:
                    lines.append("")
                    lines.extend(_table_to_markdown(tbl))
                    lines.append("")
                    break

    content = "\n".join(lines)
    content = re.sub(r'\n{3,}', '\n\n', content)
    return content.strip()


def _nsmap(prefix: str) -> str:
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    return ns.get(prefix, "")


def _get_outline_level(para) -> int | None:
    """Extract outline level from paragraph XML (w:outlineLvl)."""
    ppr = para._element.find(f'{{{_W_NS}}}pPr')
    if ppr is not None:
        ol = ppr.find(f'{{{_W_NS}}}outlineLvl')
        if ol is not None:
            try:
                return int(ol.get(f'{{{_W_NS}}}val'))
            except (TypeError, ValueError):
                pass
    return None


def _looks_like_title(text: str) -> bool:
    """Heuristic: true titles don't end with descriptive punctuation and are short."""
    if len(text) > 50:
        return False
    if text.endswith(("\u3002", "\uff1b", "\uff1a", ":", ";")):
        return False
    if re.search(r'\u53ef\u53c2\u8003|\u5e94\u5f53|\u5e94\u63d0\u4f9b|\u5e94\u660e\u786e|\u9700\u8981\u63cf\u8ff0|\u9700\u8981\u8bf4\u660e|\u9700\u8981\u5217\u8868|\u5177\u4f53\u8981\u6c42', text):
        return False
    return True


def _para_to_markdown(text: str, style_name: str, para) -> str:
    heading_styles = {
        "Heading 1", "Heading 2", "Heading 3", "Heading 4",
        "Heading 5", "Heading 6",
        "heading 1", "heading 2", "heading 3", "heading 4",
        "heading 5", "heading 6",
    }

    cn_heading_map = {}
    for i in range(1, 7):
        cn_heading_map[f"\u6807\u9898 {i}"] = i
        cn_heading_map[f"\u6807\u9898{i}"] = i

    if style_name in heading_styles:
        level_match = re.search(r'(\d)', style_name)
        level = int(level_match.group(1)) if level_match else 2
        return f"{'#' * level} {text}"

    if style_name in cn_heading_map:
        level = cn_heading_map[style_name]
        return f"{'#' * level} {text}"

    if style_name in ("Title", "title", "Subtitle", "subtitle"):
        return f"# {text}" if style_name.lower() == "title" else f"## {text}"

    outline = _get_outline_level(para)
    if outline is not None:
        if len(text) <= 30:
            level = min(outline + 1, 6)
            if level < 1:
                level = 1
            if level == 1:
                level = 2
            return f"{'#' * level} {text}"

    is_short = len(text) < 50

    if re.match(r'^[一二三四五六七八九十]+[、.．]', text) and is_short:
        return f"## {text}"
    if re.match(r'^（[一二三四五六七八九十]+）', text) and is_short:
        return f"### {text}"
    if re.match(r'^第[一二三四五六七八九十百]+[章节]', text) and is_short:
        return f"## {text}"
    if re.match(r'^第[一二三四五六七八九十百]+条', text) and is_short:
        return f"### {text}"

    m_n = re.match(r'^(\d+)[、.．]\s*(\S.*)', text, re.DOTALL)
    if m_n:
        rest_n = m_n.group(2).strip()
        if _looks_like_title(rest_n):
            return f"### {text}"

    if re.match(r'^\d+\.\d+\.\d+[.\s]', text) and is_short:
        return f"##### {text}"

    m_sub = re.match(r'^(\d+\.\d+)\s*(.*)', text, re.DOTALL)
    if m_sub:
        rest = m_sub.group(2).strip()
        if _looks_like_title(rest):
            return f"#### {text}"

    if re.match(r'^（\d+）', text) and is_short:
        return f"#### {text}"
    if re.match(r'^附[录件]', text) and is_short:
        return f"## {text}"

    is_bold = (all(run.bold for run in para.runs if run.text.strip())
               if para.runs else False)
    if is_bold and len(text) < 50:
        return f"**{text}**"

    return text


def _table_to_markdown(table) -> list[str]:
    if not table.rows:
        return []

    rows_data = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ').replace('|', '/')
                 for cell in row.cells]
        rows_data.append(cells)

    if not rows_data:
        return []

    max_cols = max(len(r) for r in rows_data)
    for row in rows_data:
        while len(row) < max_cols:
            row.append("")

    lines = []
    header = rows_data[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    for row in rows_data[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def _save_fulltext(entry: dict, markdown: str, url: str,
                   entries: list[dict]) -> bool:
    """Save validated fulltext to the entry's slug file. Returns True if index was modified."""
    title_zh = entry.get("title", {}).get("zh", "")
    doc_number = entry.get("doc_number", "")
    eid = entry.get("id", "")
    index_modified = False

    slug = entry.get("slug", "")
    if not slug:
        slug = generate_slug(title_zh, doc_number)
        existing_slugs = {e.get("slug", "") for e in entries}
        if slug in existing_slugs:
            slug = f"{slug}-{hashlib.md5(eid.encode()).hexdigest()[:4]}"
        entry["slug"] = slug
        index_modified = True
        logger.info(f"    Generated slug: {slug}")

    entry["source_url"] = url

    fulltext_md = f"# {title_zh}\n\n"
    if doc_number:
        fulltext_md += f"**{doc_number}**\n\n"
    fulltext_md += f"**Source:** [{url}]({url})\n\n"
    fulltext_md += f"---\n\n{markdown}\n"

    FULLTEXT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = FULLTEXT_DIR / f"{slug}.md"
    out_path.write_text(fulltext_md, encoding="utf-8")
    logger.info(f"    Saved: {out_path.name} ({len(fulltext_md)} chars)")
    return index_modified


def download_and_extract(limit: int = 0, dry_run: bool = False,
                         force: bool = False, category: str = ""):
    """Phase 2: Download doc/docx files, validate, extract, cross-match.

    For each entry, tries ALL candidate URLs. For each URL:
    1. Download the file
    2. Convert .doc -> .docx if needed (LibreOffice)
    3. Extract images
    4. Extract text to markdown
    5. Validate: reject catalogs
    6. If content matches searched title -> save
    7. If not -> cross-match against entire index
    8. If no cross-match -> save to _unmatched/
    """
    index_data = load_index()
    entries = index_data.get("entries", [])
    discovered = load_discovered_urls()

    spam_urls = _build_spam_url_set()
    if spam_urls:
        logger.info(f"Filtering {len(spam_urls)} high-frequency spam URLs")

    announce_url_pool = _build_announcement_url_pool(entries, discovered, spam_urls)

    candidates = []
    for entry in entries:
        eid = entry["id"]
        if category and entry.get("category") != category:
            continue
        if not force and fulltext_exists(entry):
            continue
        disc = discovered.get(eid, {})
        urls = disc.get("urls", [])
        if not urls:
            old_url = disc.get("url", "")
            if old_url:
                urls = [old_url]
        urls = [u for u in urls if u not in spam_urls]

        sibling_urls = announce_url_pool.get(eid, set()) - set(urls)
        if sibling_urls:
            urls = urls + sorted(sibling_urls)

        if not urls:
            continue
        candidates.append((entry, urls))

    if limit > 0:
        candidates = candidates[:limit]

    logger.info(f"Processing {len(candidates)} entries")

    success = 0
    cross_matched = 0
    unmatched_saved = 0
    failed = 0
    index_modified = False

    FULLTEXT_DIR.mkdir(parents=True, exist_ok=True)

    seen_urls_this_run: set[str] = set()

    for i, (entry, urls) in enumerate(candidates):
        title_zh = entry.get("title", {}).get("zh", "")
        eid = entry["id"]

        logger.info(f"[{i+1}/{len(candidates)}] {title_zh[:55]}")

        if dry_run:
            continue

        entry_saved = False

        for j, url in enumerate(urls):
            if url in seen_urls_this_run:
                continue

            logger.info(f"  Try URL {j+1}/{len(urls)}: ...{url[-50:]}")

            raw_bytes = download_file(url)
            if not raw_bytes:
                continue

            size_kb = len(raw_bytes) / 1024
            logger.info(f"    Downloaded {size_kb:.1f} KB")

            docx_bytes = file_to_docx_bytes(raw_bytes, url)
            if not docx_bytes:
                continue

            slug_for_images = entry.get("slug", "") or generate_slug(
                title_zh, entry.get("doc_number", ""))

            image_map = extract_images_from_docx(docx_bytes, slug_for_images)

            markdown = extract_docx_to_markdown(docx_bytes, image_map)
            if not markdown or len(markdown) < 100:
                logger.info(f"    Skip: extraction failed or too short")
                continue

            seen_urls_this_run.add(url)

            if is_catalog_content(markdown):
                logger.info(f"    REJECTED: catalog")
                if image_map:
                    img_dir = IMAGES_DIR / slug_for_images
                    if img_dir.exists():
                        shutil.rmtree(img_dir)
                continue

            if is_draft_content(markdown):
                logger.info(f"    REJECTED: draft for comments (征求意见稿)")
                if image_map:
                    img_dir = IMAGES_DIR / slug_for_images
                    if img_dir.exists():
                        shutil.rmtree(img_dir)
                continue

            if content_matches_title(markdown, title_zh):
                if _save_fulltext(entry, markdown, url, entries):
                    index_modified = True
                logger.info(f"    MATCHED to searched entry")
                success += 1
                entry_saved = True
                break

            logger.info(f"    Title mismatch -- trying cross-match...")
            matched_entry = cross_match_content(markdown, entries, url)
            if matched_entry:
                matched_title = matched_entry["title"]["zh"]
                logger.info(f"    CROSS-MATCHED: {matched_title[:50]}")

                cm_slug = matched_entry.get("slug", "") or generate_slug(
                    matched_title, matched_entry.get("doc_number", ""))
                if image_map:
                    old_dir = IMAGES_DIR / slug_for_images
                    new_dir = IMAGES_DIR / cm_slug
                    if old_dir.exists() and slug_for_images != cm_slug:
                        if new_dir.exists():
                            shutil.rmtree(new_dir)
                        old_dir.rename(new_dir)
                        new_map = {}
                        for rid, path in image_map.items():
                            new_map[rid] = path.replace(
                                f"/images/nmpa-guidance/{slug_for_images}/",
                                f"/images/nmpa-guidance/{cm_slug}/")
                        image_map = new_map
                        markdown = extract_docx_to_markdown(docx_bytes, image_map)

                if _save_fulltext(matched_entry, markdown, url, entries):
                    index_modified = True
                cross_matched += 1
                entry_saved = True
                break

            actual_title = extract_actual_title(markdown)
            save_unmatched(markdown, url, actual_title)
            unmatched_saved += 1
            if image_map:
                img_dir = IMAGES_DIR / slug_for_images
                if img_dir.exists():
                    shutil.rmtree(img_dir)

            time.sleep(0.3)

        if not entry_saved and not dry_run:
            failed += 1

        time.sleep(0.3)

    if index_modified and not dry_run:
        save_index(index_data)

    logger.info(
        f"\nResults: {success} direct + {cross_matched} cross-matched "
        f"+ {unmatched_saved} unmatched, {failed} failed"
    )


# ---------------------------------------------------------------------------
# Cleanup: remove bad fulltext files
# ---------------------------------------------------------------------------

def cleanup_bad_fulltext(category: str = "", dry_run: bool = False,
                         protect_slugs: set[str] | None = None):
    """Remove fulltext files that are catalogs/lists rather than real content."""
    index_data = load_index()
    entries = index_data.get("entries", [])

    removed = 0
    checked = 0
    index_modified = False

    for entry in entries:
        slug = entry.get("slug", "")
        if not slug:
            continue
        if category and entry.get("category") != category:
            continue

        fp = FULLTEXT_DIR / f"{slug}.md"
        if not fp.exists():
            continue

        checked += 1

        if protect_slugs and slug in protect_slugs:
            continue

        content = fp.read_text(encoding="utf-8")

        body = content
        separator_pos = content.find("\n---\n")
        if separator_pos >= 0:
            body = content[separator_pos + 5:]

        title_zh = entry.get("title", {}).get("zh", "")
        is_valid, reason = validate_extracted_content(body, title_zh)

        if not is_valid:
            logger.info(
                f"BAD [{reason}]: {title_zh[:50]} -> {slug}.md"
            )
            if not dry_run:
                fp.unlink()
                en_fp = FULLTEXT_DIR / f"{slug}.en.md"
                if en_fp.exists():
                    en_fp.unlink()
                    logger.info(f"  Also removed: {slug}.en.md")
                img_dir = IMAGES_DIR / slug
                if img_dir.exists():
                    shutil.rmtree(img_dir)
                    logger.info(f"  Also removed images: {slug}/")
                entry.pop("slug", None)
                entry.pop("source_url", None)
                index_modified = True
            removed += 1

    if index_modified and not dry_run:
        save_index(index_data)

    logger.info(
        f"\nCleanup: checked {checked}, removed {removed}"
        f"{' (dry-run)' if dry_run else ''}"
    )


# ---------------------------------------------------------------------------
# Stats
# ---------------------------------------------------------------------------

def show_stats():
    index_data = load_index()
    entries = index_data.get("entries", [])
    discovered = load_discovered_urls()

    with_slug = [e for e in entries if e.get("slug")]
    with_fulltext = [e for e in entries if fulltext_exists(e)]
    without_fulltext = [e for e in entries if not fulltext_exists(e)]
    discovered_with_url = {
        k: v for k, v in discovered.items()
        if v.get("urls") or v.get("url")
    }
    discovered_no_url = {
        k: v for k, v in discovered.items()
        if not v.get("urls") and not v.get("url")
    }

    ft_files = list(FULLTEXT_DIR.glob("*.md")) if FULLTEXT_DIR.exists() else []
    en_files = [f for f in ft_files if f.stem.endswith('.en')]
    zh_only = [f for f in ft_files if not f.stem.endswith('.en')]

    unmatched_files = list(UNMATCHED_DIR.glob("*.md")) if UNMATCHED_DIR.exists() else []

    lo_bin = _detect_libreoffice()

    print(f"\n=== NMPA Guidance Fulltext Statistics ===")
    print(f"Total index entries:       {len(entries)}")
    print(f"Entries with slug:         {len(with_slug)}")
    print(f"Entries with fulltext:     {len(with_fulltext)}")
    print(f"Entries without fulltext:  {len(without_fulltext)}")
    print(f"")
    print(f"Fulltext files (ZH *.md):  {len(zh_only)}")
    print(f"English translations:      {len(en_files)}")
    print(f"Unmatched files:           {len(unmatched_files)}")
    print(f"")
    print(f"LibreOffice (.doc support):{' YES' if lo_bin else ' NO'}")
    print(f"")
    print(f"URL discovery status:")
    print(f"  Searched:                {len(discovered)}")
    print(f"  Found doc/docx URL:      {len(discovered_with_url)}")
    print(f"  No URL found:            {len(discovered_no_url)}")
    not_searched = len(without_fulltext) - len(
        [e for e in without_fulltext if e["id"] in discovered]
    )
    print(f"  Not yet searched:        {not_searched}")
    print()

    by_cat = {}
    for e in entries:
        cat = e.get("category", "unknown")
        ft = fulltext_exists(e)
        by_cat.setdefault(cat, {"total": 0, "has_ft": 0})
        by_cat[cat]["total"] += 1
        if ft:
            by_cat[cat]["has_ft"] += 1

    print("Coverage by category:")
    for cat, s in sorted(by_cat.items(), key=lambda x: -x[1]["total"]):
        pct = s["has_ft"] / s["total"] * 100 if s["total"] else 0
        print(f"  {cat:25s} {s['has_ft']:3d}/{s['total']:3d} ({pct:.0f}%)")


def main():
    parser = argparse.ArgumentParser(
        description="Fetch NMPA guidance fulltext documents",
    )
    parser.add_argument("--discover", action="store_true",
                        help="Phase 1: Search for doc/docx download URLs")
    parser.add_argument("--download", action="store_true",
                        help="Phase 2: Download, convert, validate, extract")
    parser.add_argument("--cleanup", action="store_true",
                        help="Remove bad fulltext (catalogs/lists)")
    parser.add_argument("--rematch", action="store_true",
                        help="Re-match _unmatched/ files to index entries")
    parser.add_argument("--stats", action="store_true",
                        help="Show fulltext coverage statistics")
    parser.add_argument("--limit", type=int, default=0,
                        help="Max entries to process (0 = all)")
    parser.add_argument("--dry-run", action="store_true",
                        help="Preview without downloading/writing")
    parser.add_argument("--force", action="store_true",
                        help="Re-process even if fulltext exists")
    parser.add_argument("--category", type=str, default="",
                        help="Filter by category (e.g. general, ivd)")
    args = parser.parse_args()

    if not (args.discover or args.download or args.cleanup
            or args.rematch or args.stats):
        parser.print_help()
        print("\nExamples:")
        print("  python scripts/fetch_nmpa_fulltext.py --stats")
        print("  python scripts/fetch_nmpa_fulltext.py --cleanup --category general")
        print("  python scripts/fetch_nmpa_fulltext.py --discover --category general")
        print("  python scripts/fetch_nmpa_fulltext.py --download --category general")
        print("  python scripts/fetch_nmpa_fulltext.py --rematch")
        return

    if args.stats:
        show_stats()
        return

    if args.cleanup:
        cleanup_bad_fulltext(category=args.category, dry_run=args.dry_run)
        return

    if args.rematch:
        rematch_unmatched(dry_run=args.dry_run)
        return

    searcher = WebSearcher()

    if args.discover:
        discover_urls(searcher, limit=args.limit, dry_run=args.dry_run,
                      category=args.category)

    if args.download:
        download_and_extract(limit=args.limit, dry_run=args.dry_run,
                             force=args.force, category=args.category)


if __name__ == "__main__":
    main()
